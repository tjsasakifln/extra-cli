"""Traceable extraction for PDF, DOCX, XLSX, HTML, TXT."""

from __future__ import annotations

import re
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from scripts.edital_case.store import get_object_path, write_json


def _jsonl_write(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as fh:
        import json

        for row in rows:
            fh.write(json.dumps(row, ensure_ascii=False, default=str) + "\n")


def extract_pdf(path: Path, document_id: str) -> dict[str, Any]:
    blocks: list[dict[str, Any]] = []
    method = "none"
    quality = "EMPTY"
    page_count = 0
    errors: list[str] = []
    try:
        import fitz  # pymupdf

        doc = fitz.open(str(path))
        page_count = doc.page_count
        method = "pymupdf_text"
        for i in range(page_count):
            page = doc.load_page(i)
            text = page.get_text("text") or ""
            text = text.replace("\x00", "")
            status = "OK" if text.strip() else "OCR_REQUIRED"
            blocks.append(
                {
                    "document_id": document_id,
                    "page": i + 1,
                    "text": text,
                    "extraction_method": method,
                    "char_count": len(text),
                    "quality_status": status,
                    "locator": f"page:{i + 1}",
                }
            )
        doc.close()
    except Exception as exc:  # noqa: BLE001
        errors.append(f"pymupdf: {exc}")
        # fallback pypdf then pdfminer
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            method = "pypdf"
            page_count = len(reader.pages)
            for i, page in enumerate(reader.pages, start=1):
                text = (page.extract_text() or "").replace("\x00", "")
                blocks.append(
                    {
                        "document_id": document_id,
                        "page": i,
                        "text": text,
                        "extraction_method": method,
                        "char_count": len(text),
                        "quality_status": "OK" if text.strip() else "OCR_REQUIRED",
                        "locator": f"page:{i}",
                    }
                )
        except Exception as exc_pypdf:  # noqa: BLE001
            errors.append(f"pypdf: {exc_pypdf}")
            try:
                from pdfminer.high_level import extract_pages
                from pdfminer.layout import LTTextContainer

                method = "pdfminer"
                for i, layout in enumerate(extract_pages(str(path)), start=1):
                    parts: list[str] = []
                    for element in layout:
                        if isinstance(element, LTTextContainer):
                            parts.append(element.get_text())
                    text = "".join(parts)
                    blocks.append(
                        {
                            "document_id": document_id,
                            "page": i,
                            "text": text,
                            "extraction_method": method,
                            "char_count": len(text),
                            "quality_status": "OK" if text.strip() else "OCR_REQUIRED",
                            "locator": f"page:{i}",
                        }
                    )
                page_count = len(blocks)
            except Exception as exc2:  # noqa: BLE001
                errors.append(f"pdfminer: {exc2}")
                return {
                    "document_id": document_id,
                    "format": "pdf",
                    "status": "EXTRACTION_FAILED",
                    "error": "; ".join(errors),
                    "blocks": [],
                    "page_count": 0,
                    "total_chars": 0,
                    "quality_status": "EXTRACTION_FAILED",
                    "extraction_method": method,
                }

    total_chars = sum(b["char_count"] for b in blocks)
    ocr_pages = sum(1 for b in blocks if b["quality_status"] == "OCR_REQUIRED")
    if total_chars == 0:
        quality = "OCR_REQUIRED"
    elif ocr_pages and ocr_pages == len(blocks):
        quality = "OCR_REQUIRED"
    elif ocr_pages:
        quality = "PARTIAL"
    else:
        quality = "OK"

    return {
        "document_id": document_id,
        "format": "pdf",
        "status": "OK" if total_chars or blocks else "EMPTY",
        "blocks": blocks,
        "page_count": page_count or len(blocks),
        "total_chars": total_chars,
        "quality_status": quality,
        "extraction_method": method,
        "ocr_used": False,
        "ocr_pages": ocr_pages,
    }


def extract_docx(path: Path, document_id: str) -> dict[str, Any]:
    blocks: list[dict[str, Any]] = []
    tables_out: list[dict[str, Any]] = []
    try:
        from docx import Document

        doc = Document(str(path))
        for i, para in enumerate(doc.paragraphs):
            text = para.text or ""
            style = getattr(para.style, "name", None) if para.style else None
            is_heading = bool(style and "Heading" in style)
            blocks.append(
                {
                    "document_id": document_id,
                    "page": None,
                    "paragraph": i,
                    "text": text,
                    "extraction_method": "python-docx",
                    "char_count": len(text),
                    "quality_status": "OK" if text.strip() else "EMPTY",
                    "locator": f"paragraph:{i}",
                    "style": style,
                    "is_heading": is_heading,
                }
            )
        for ti, table in enumerate(doc.tables):
            rows_data = []
            for ri, row in enumerate(table.rows):
                cells = []
                for ci, cell in enumerate(row.cells):
                    cells.append(
                        {
                            "row": ri,
                            "col": ci,
                            "text": cell.text or "",
                            "locator": f"table:{ti}:cell:{ri}:{ci}",
                        }
                    )
                rows_data.append(cells)
            tables_out.append({"table_index": ti, "rows": rows_data})
    except Exception as exc:  # noqa: BLE001
        return {
            "document_id": document_id,
            "format": "docx",
            "status": "EXTRACTION_FAILED",
            "error": str(exc),
            "blocks": [],
            "tables": [],
            "total_chars": 0,
            "quality_status": "EXTRACTION_FAILED",
            "extraction_method": "python-docx",
        }
    total = sum(b["char_count"] for b in blocks)
    return {
        "document_id": document_id,
        "format": "docx",
        "status": "OK",
        "blocks": blocks,
        "tables": tables_out,
        "total_chars": total,
        "quality_status": "OK" if total else "EMPTY",
        "extraction_method": "python-docx",
        "page_count": None,
    }


def extract_xlsx(path: Path, document_id: str) -> dict[str, Any]:
    tables_out: list[dict[str, Any]] = []
    blocks: list[dict[str, Any]] = []
    try:
        import openpyxl

        wb = openpyxl.load_workbook(str(path), data_only=False)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            hidden = ws.sheet_state != "visible"
            cells = []
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is None:
                        continue
                    coord = cell.coordinate
                    formula = None
                    value = cell.value
                    if isinstance(value, str) and value.startswith("="):
                        formula = value
                    locator = f"sheet:{sheet_name}!{coord}"
                    cell_rec = {
                        "sheet": sheet_name,
                        "coordinate": coord,
                        "value": value if not formula else None,
                        "formula": formula,
                        "locator": locator,
                        "hidden_sheet": hidden,
                    }
                    cells.append(cell_rec)
                    text = formula or (str(value) if value is not None else "")
                    blocks.append(
                        {
                            "document_id": document_id,
                            "page": None,
                            "text": text,
                            "extraction_method": "openpyxl",
                            "char_count": len(text),
                            "quality_status": "OK",
                            "locator": locator,
                            "sheet": sheet_name,
                            "cell": coord,
                        }
                    )
            merged = [str(m) for m in ws.merged_cells.ranges]
            tables_out.append(
                {
                    "sheet": sheet_name,
                    "hidden": hidden,
                    "dimensions": ws.dimensions,
                    "merged_ranges": merged,
                    "cells": cells[:5000],  # cap
                    "cell_count": len(cells),
                }
            )
    except Exception as exc:  # noqa: BLE001
        return {
            "document_id": document_id,
            "format": "xlsx",
            "status": "EXTRACTION_FAILED",
            "error": str(exc),
            "blocks": [],
            "tables": [],
            "total_chars": 0,
            "quality_status": "EXTRACTION_FAILED",
            "extraction_method": "openpyxl",
        }
    total = sum(b["char_count"] for b in blocks)
    return {
        "document_id": document_id,
        "format": "xlsx",
        "status": "OK",
        "blocks": blocks,
        "tables": tables_out,
        "total_chars": total,
        "quality_status": "OK" if total else "EMPTY",
        "extraction_method": "openpyxl",
        "page_count": None,
    }


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[tuple[str, str]] = []
        self._skip = False
        self._order = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"script", "style"}:
            self._skip = True
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td", "th", "tr", "div"}:
            self._order += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style"}:
            self._skip = False

    def handle_data(self, data: str) -> None:
        if self._skip:
            return
        text = data.strip()
        if text:
            self.parts.append((f"html:block:{self._order}", text))


def extract_html(path: Path, document_id: str) -> dict[str, Any]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    parser = _HTMLTextExtractor()
    try:
        parser.feed(raw)
    except Exception as exc:  # noqa: BLE001
        return {
            "document_id": document_id,
            "format": "html",
            "status": "EXTRACTION_FAILED",
            "error": str(exc),
            "blocks": [],
            "total_chars": 0,
            "quality_status": "EXTRACTION_FAILED",
            "extraction_method": "html.parser",
        }
    blocks = []
    for i, (loc, text) in enumerate(parser.parts):
        blocks.append(
            {
                "document_id": document_id,
                "page": None,
                "text": text,
                "extraction_method": "html.parser",
                "char_count": len(text),
                "quality_status": "OK",
                "locator": loc or f"html:block:{i}",
            }
        )
    total = sum(b["char_count"] for b in blocks)
    return {
        "document_id": document_id,
        "format": "html",
        "status": "OK",
        "blocks": blocks,
        "total_chars": total,
        "quality_status": "OK" if total else "EMPTY",
        "extraction_method": "html.parser",
        "page_count": None,
    }


def extract_txt(path: Path, document_id: str) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8", errors="replace")
    lines = text.splitlines()
    blocks = []
    for i, line in enumerate(lines):
        if not line.strip():
            continue
        blocks.append(
            {
                "document_id": document_id,
                "page": None,
                "text": line,
                "extraction_method": "text",
                "char_count": len(line),
                "quality_status": "OK",
                "locator": f"line:{i + 1}",
            }
        )
    return {
        "document_id": document_id,
        "format": "txt",
        "status": "OK",
        "blocks": blocks,
        "total_chars": len(text),
        "quality_status": "OK" if text.strip() else "EMPTY",
        "extraction_method": "text",
        "page_count": None,
    }


def extract_document(
    case_dir: Path,
    *,
    document_id: str,
    sha256: str,
    extension: str,
    original_name: str,
) -> dict[str, Any]:
    path = get_object_path(case_dir, sha256)
    ext = extension.lower()
    if ext == ".pdf":
        result = extract_pdf(path, document_id)
    elif ext == ".docx":
        result = extract_docx(path, document_id)
    elif ext in {".xlsx", ".xlsm"}:
        result = extract_xlsx(path, document_id)
    elif ext in {".html", ".htm"}:
        result = extract_html(path, document_id)
    elif ext in {".txt", ".md", ".csv", ".json", ".xml"}:
        result = extract_txt(path, document_id)
    else:
        result = {
            "document_id": document_id,
            "format": ext.lstrip(".") or "unknown",
            "status": "UNSUPPORTED",
            "blocks": [],
            "total_chars": 0,
            "quality_status": "UNSUPPORTED",
            "extraction_method": "none",
            "page_count": None,
        }

    doc_dir = case_dir / "documents" / document_id
    doc_dir.mkdir(parents=True, exist_ok=True)
    blocks = result.get("blocks") or []
    _jsonl_write(doc_dir / "extraction.jsonl", blocks)
    write_json(doc_dir / "tables.json", result.get("tables") or [])
    # references filled later
    if not (doc_dir / "references.json").exists():
        write_json(doc_dir / "references.json", {"references": []})

    # compact extraction summary without full text duplication in document.json later
    summary = {k: v for k, v in result.items() if k not in {"blocks", "tables"}}
    summary["block_count"] = len(blocks)
    summary["original_name"] = original_name
    summary["sha256"] = sha256
    write_json(doc_dir / "extraction-summary.json", summary)
    return result


def load_extraction_blocks(case_dir: Path, document_id: str) -> list[dict[str, Any]]:
    path = case_dir / "documents" / document_id / "extraction.jsonl"
    if not path.exists():
        return []
    import json

    rows = []
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.strip():
            rows.append(json.loads(line))
    return rows


def full_text(blocks: list[dict[str, Any]]) -> str:
    return "\n".join(b.get("text") or "" for b in blocks)


def find_excerpt(blocks: list[dict[str, Any]], pattern: str, *, flags: int = re.I) -> dict[str, Any] | None:
    rx = re.compile(pattern, flags)
    for b in blocks:
        text = b.get("text") or ""
        m = rx.search(text)
        if not m:
            continue
        start = max(0, m.start() - 80)
        end = min(len(text), m.end() + 120)
        # Snap start to a token boundary so citation excerpts are verifiable
        # substrings (avoid mid-word windows like "RÔNICO" from "ELETRÔNICO").
        if start > 0 and text[start].isalnum() and text[start - 1].isalnum():
            boundary = max(text.rfind(" ", 0, m.start()), text.rfind("\n", 0, m.start()))
            if boundary >= max(0, m.start() - 80):
                start = boundary + 1
            else:
                start = m.start()
        return {
            "document_id": b.get("document_id"),
            "page": b.get("page"),
            "locator": b.get("locator"),
            "paragraph": b.get("paragraph"),
            "cell": b.get("cell"),
            "excerpt": text[start:end].strip(),
            "match": m.group(0),
        }
    return None
